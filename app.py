import streamlit as st
import anthropic
import json
import io
import re

# ─── Page Configuration ───────────────────────────────────────────────────────
st.set_page_config(
    page_title="Term Sheet Scanner",
    page_icon="📄",
    layout="wide"
)

# ─── Custom CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .main-header { text-align: center; padding: 1.5rem 0 0.5rem 0; }
    .main-header h1 { font-size: 2.2rem; font-weight: 700; color: #1a1a2e; margin-bottom: 0.3rem; }
    .main-header p { color: #666; font-size: 1rem; }
    .summary-box {
        background-color: #f8f9fa;
        border-left: 4px solid #185FA5;
        border-radius: 4px;
        padding: 1rem 1.25rem;
        margin-bottom: 1.5rem;
    }
    .disclaimer {
        text-align: center; color: #999; font-size: 0.8rem;
        margin-top: 2rem; padding-top: 1rem; border-top: 1px solid #eee;
    }
    div[data-testid="stFileUploader"] { border: 2px dashed #185FA5; border-radius: 8px; padding: 1rem; }

    /* ── Compact table styles ── */
    .compact-table-wrapper {
        overflow-x: auto;
        overflow-y: auto;
        max-height: 520px;
        border: 1px solid #d0d0d0;
        border-radius: 6px;
        margin-bottom: 1rem;
    }
    .compact-table {
        border-collapse: collapse;
        width: 100%;
        font-size: 12.5px;
        font-family: sans-serif;
        white-space: nowrap;
    }
    .compact-table thead th {
        background-color: #185FA5;
        color: white;
        padding: 8px 12px;
        text-align: left;
        font-weight: 600;
        position: sticky;
        top: 0;
        z-index: 2;
        white-space: nowrap;
    }
    .compact-table tbody tr { border-bottom: 1px solid #e8e8e8; }
    .compact-table tbody tr:hover { filter: brightness(0.96); }
    .compact-table tbody td {
        padding: 6px 12px;
        vertical-align: middle;
        max-width: 220px;
        overflow: hidden;
        text-overflow: ellipsis;
        white-space: nowrap;
        cursor: default;
    }
    .compact-table tbody td:first-child { font-weight: 600; min-width: 130px; }
    .row-nn { background-color: #fff0f0; }
    .row-cl { background-color: #fffbf0; }
    .row-gh { background-color: #f0fff0; }
    .badge {
        display: inline-block; font-size: 10.5px; padding: 2px 8px;
        border-radius: 12px; font-weight: 600; white-space: nowrap;
    }
    .badge-nn { background: #FCEBEB; color: #A32D2D; }
    .badge-cl { background: #FAEEDA; color: #854F0B; }
    .badge-gh { background: #EAF3DE; color: #3B6D11; }
</style>
""", unsafe_allow_html=True)

# ─── System Prompt ────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are a seasoned Debt Capital Markets expert with 15+ years of experience \
raising debt for NBFCs across all stages of growth — from early-stage NBFCs with Rs 100-300 Cr \
AUM to large established NBFCs with Rs 5,000+ Cr AUM. You have deep experience across all \
debt instruments: bilateral term loans from banks and NBFCs, Non-Convertible Debentures (NCDs) \
both listed and unlisted, Pass-Through Certificates (PTCs), Direct Assignments (DAs), and \
co-lending arrangements.

Your analysis is always calibrated to: (a) the size and growth stage of the borrowing NBFC \
as evident from the term sheet, (b) the type of instrument being offered, and (c) what is \
genuinely standard market practice for that specific instrument and lender type.

---

STEP 1: If the document is NOT a term sheet or lending agreement, respond ONLY with: \
{"not_termsheet": true}

---

STEP 2: Before analysing, identify from the term sheet:
- Instrument type: BILATERAL_BANK_LOAN / BILATERAL_NBFC_LOAN / NCD_LISTED / NCD_UNLISTED \
/ PTC / DIRECT_ASSIGNMENT / CO_LENDING / OTHER
- Borrower size: infer from AUM, facility size, or borrower profile mentioned
- Lender type: Bank / NBFC / AIF / Debt Capital Markets / Multiple Investors / Other

Calibrate your entire analysis based on these three factors.

---

STEP 3: Analyse using the priority framework below.

## TIER 1 — ALWAYS ANALYSE (financial and operational impact):

1. Financial Covenants — flag if: measured at standalone NBFC level not consolidated \
(always push for consolidated including managed AUM); thresholds tight for NBFC size/stage; \
no cure period before enforcement; definitions of TNW/leverage/NPA are narrow.

2. Negative Covenants — flag if: additional borrowing restriction tighter than the financial \
covenant leverage ratio already provides; no carve-outs for co-lending or securitisation; \
related party restrictions cover ordinary course business; dividend restrictions apply even \
when all covenants are met.

3. Events of Default — flag if: cure period for payment default below 5 business days; \
cure period for covenant breach below 30 days; cross-default threshold below Rs 1 Cr; \
MAC definition entirely subjective; KMP exit is automatic EoD; cross-default triggered \
by non-financial defaults of group entities.

4. Prepayment Terms — flag if: lock-in period exceeds 12 months; prepayment penalty exceeds \
1% after lock-in; no right to prepay penalty-free upon lender assignment to unacceptable \
assignee or upon unilateral repricing.

5. Security and Collateral — flag if: personal guarantees of promoter/directors required \
(not market standard for institutional borrowing above Rs 25 Cr); security cover exceeds \
1.1x for receivables-backed; asset replacement required within less than 30 days; \
quarterly charge registration amendment required (push for annual).

6. Default Interest Rate — flag if: exceeds 2% per annum over base rate; applies from \
date of default without any cure period notice.

7. Discretionary Audits — flag if: more frequent than semi-annual; notice period less \
than 5 business days; scope beyond credit, collections, and portfolio quality.

8. Assignment Rights — flag if: no restriction on assignee type including competitors \
or distressed asset buyers; notice period less than 7 business days; no penalty-free \
prepayment right on unacceptable assignment.

9. Reporting Covenants — flag ONLY if: unaudited financials required within less than \
45 days; audited financials within less than 120 days; timelines operationally unrealistic.

10. Tenor and Repayment — flag if: repayment creates ALM mismatch with NBFC loan book \
duration; availability period less than 30 days; bullet repayment on tenor shorter than \
weighted average asset duration.

FOR NCD INSTRUMENTS — additionally flag:
- Rating trigger clauses causing acceleration or coupon step-up above 50 bps per notch
- Asset cover ratio above 1.1x or CA certificate more frequent than semi-annual
- Put options exercisable on rating events or covenant breach
- Debenture trustee unilateral EoD powers without bondholder resolution (75% threshold)

FOR PTC / DIRECT ASSIGNMENT — additionally flag:
- Eligibility criteria more restrictive than NBFC's own underwriting policy
- FLDG above 5% of pool; substitution rights restricted post pool cutoff
- Servicer termination without cause or less than 30 days notice

FOR CO-LENDING — additionally flag:
- Co-lending ratio changeable unilaterally by bank partner
- DLG/FLDG settlement delayed beyond 30 days
- Bank credit policy overriding NBFC underwriting without consent
- Exit with less than 90 days notice without penalty

## TIER 2 — DO NOT FLAG (industry standard — accepted market practice):
- Unilateral lender discretion to amend sanction letter pre-execution (industry norm)
- Non-binding / no-commitment language in sanction letter (industry norm)
- Standard representations and warranties (corporate existence, authority, no litigation)
- Standard conditions precedent (board resolutions, constitutional docs, due diligence)
- Standard affirmative covenants (comply with law, pay taxes, maintain licences)
- Governing law and jurisdiction
- Business day convention and definitions
- Costs and expenses borne by borrower on bilateral loans
- Finance document composition
- For NCDs: SEBI ILDS/OBPP compliance, standard trustee appointment, SEBI-mandated \
disclosure timelines (statutory — cannot be negotiated)

## SIZE CALIBRATION:
Small NBFC (below Rs 500 Cr AUM): Focus on most critical terms only — financial covenants, \
EoD cure periods, personal guarantees. Accept some unfavourable terms as cost of institutional \
access. Mark fewer items as Non-negotiable.
Mid-sized (Rs 500 Cr – Rs 2,000 Cr): Full push on all Tier 1 clauses.
Large (above Rs 2,000 Cr): Flag even moderately unfavourable terms as Non-negotiable.

## CRITICALITY:
Non-negotiable: Must change — meaningful breach risk OR severe consequence OR material \
growth constraint.
Can live with it: Low breach probability or immaterial consequence — propose but do not \
push hard if lender resists.
Good to have: Concede gracefully if lender declines.

---

STEP 4: Respond ONLY with valid JSON — no markdown, no backticks, no extra text:

{
  "instrument_type": "BILATERAL_BANK_LOAN",
  "borrower_profile": "One sentence on NBFC size and stage inferred from the term sheet",
  "summary": {
    "overall": "One sentence overall assessment",
    "non_negotiable_count": 0,
    "most_critical": "Name of most critical clause"
  },
  "clauses": [
    {
      "clause_name": "Short clause name",
      "clause_details": "Exact or near-exact clause language from the term sheet",
      "implications": "Specific problem this creates. Note if tighter/in line/more relaxed \
vs market standard for this instrument type.",
      "criticality": "Non-negotiable",
      "lenders_pushback": "Their likely counter-argument and how to respond to it.",
      "revised_clause": "Exact revised clause language to propose to the lender.",
      "rationale": "Why this change is needed, how it addresses the constraint, and how \
lender risk is still protected."
    }
  ]
}

Sort: Non-negotiable first, then Can live with it, then Good to have.
Within each tier: financial materiality — most impactful first.
Maximum 12 clauses. Tier 1 types only — include Tier 2 only if genuinely outside market norms.
CRITICAL: Each JSON field must not exceed 50 words. Be ruthlessly concise. No generic observations."""

# ─── Text Extraction Functions ────────────────────────────────────────────────
def extract_from_pdf(file_bytes):
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
        return text.strip()
    except ImportError:
        st.error("PyMuPDF not installed. Add 'pymupdf' to requirements.txt")
        return None
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return None

def extract_from_docx(file_bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        full_text = []

        # Extract standalone paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Extract text from ALL tables — critical for term sheets formatted as tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_content = cell.text.strip()
                    if cell_content:
                        row_text.append(cell_content)
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except ImportError:
        st.error("python-docx not installed. Add 'python-docx' to requirements.txt")
        return None
    except Exception as e:
        st.error(f"Error reading Word doc: {e}")
        return None

def extract_from_excel(file_bytes):
    try:
        import pandas as pd
        dfs = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        return "\n".join(f"\n--- Sheet: {k} ---\n{v.to_string(index=False)}" for k, v in dfs.items())
    except Exception as e:
        st.error(f"Error reading Excel: {e}")
        return None

def extract_from_image(file_bytes, filename):
    try:
        import base64
        ext = filename.split('.')[-1].lower()
        media_type = {'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png'}.get(ext, 'image/jpeg')
        return {"type": "image", "data": base64.standard_b64encode(file_bytes).decode("utf-8"), "media_type": media_type}
    except Exception as e:
        st.error(f"Error processing image: {e}")
        return None

def extract_text(uploaded_file):
    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()
    if filename.endswith('.pdf'):
        return extract_from_pdf(file_bytes), "text"
    elif filename.endswith(('.docx', '.doc')):
        return extract_from_docx(file_bytes), "text"
    elif filename.endswith(('.xlsx', '.xls')):
        return extract_from_excel(file_bytes), "text"
    elif filename.endswith(('.jpg', '.jpeg', '.png')):
        return extract_from_image(file_bytes, filename), "image"
    return None, "unsupported"

# ─── Claude API Call ──────────────────────────────────────────────────────────
def analyse_term_sheet(content, content_type, api_key):
    try:
        client = anthropic.Anthropic(api_key=api_key)
        if content_type == "image":
            user_content = [
                {"type": "image", "source": {"type": "base64", "media_type": content["media_type"], "data": content["data"]}},
                {"type": "text", "text": "Please analyse this term sheet image according to your instructions."}
            ]
        else:
            truncated = content[:6000] + ("\n\n[Document truncated]" if len(content) > 6000 else "")
            user_content = f"Please analyse this term sheet:\n\n{truncated}"

        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8000,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_content}]
        )
        raw = message.content[0].text.strip()
        clean = re.sub(r'```json|```', '', raw).strip()

        # ── Capture token usage and compute cost ──
        input_tokens  = message.usage.input_tokens
        output_tokens = message.usage.output_tokens

        # Claude Sonnet 4.6 pricing: $3/M input, $15/M output
        input_cost  = (input_tokens  / 1_000_000) * 3.00
        output_cost = (output_tokens / 1_000_000) * 15.00
        total_cost  = input_cost + output_cost

        usage = {
            "input_tokens":  input_tokens,
            "output_tokens": output_tokens,
            "input_cost_usd":  round(input_cost,  6),
            "output_cost_usd": round(output_cost, 6),
            "total_cost_usd":  round(total_cost,  6),
            "total_cost_inr":  round(total_cost * 84, 4),   # approx USD→INR
        }

        return json.loads(clean), None, usage
    except json.JSONDecodeError as e:
        return None, f"Analysis could not be processed (JSON parse error). Please try again. Detail: {e}", None
    except anthropic.AuthenticationError:
        return None, "Invalid API key. Please check your Anthropic API key in Streamlit secrets.", None
    except anthropic.RateLimitError:
        return None, "Rate limit reached. Please wait a moment and try again.", None
    except Exception as e:
        return None, f"Analysis failed: {e}", None

# ─── Display Functions ────────────────────────────────────────────────────────
def display_summary(summary):
    st.markdown(f"""
    <div class="summary-box">
        <strong>Overall Assessment:</strong> {summary.get('overall','N/A')}<br>
        <strong>Non-negotiable changes required:</strong> {summary.get('non_negotiable_count',0)}<br>
        <strong>Most critical clause:</strong> {summary.get('most_critical','N/A')}
    </div>""", unsafe_allow_html=True)

def display_compact_table(clauses):
    """Render a compact, single-line-per-row HTML table like Image 2."""
    def badge(c):
        if c == "Non-negotiable":
            return '<span class="badge badge-nn">🔴 Non-negotiable</span>'
        elif c == "Can live with it":
            return '<span class="badge badge-cl">🟡 Can live with it</span>'
        return '<span class="badge badge-gh">🟢 Good to have</span>'

    def row_class(c):
        return {"Non-negotiable": "row-nn", "Can live with it": "row-cl", "Good to have": "row-gh"}.get(c, "")

    def esc(s):
        return str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace('"','&quot;')

    headers = ["Clause Name","Clause Details","Implications","Criticality",
               "Lender's Likely Pushback","Revised Clause","Rationale for Modification"]

    rows_html = ""
    for c in clauses:
        vals = [
            esc(c.get("clause_name","")),
            esc(c.get("clause_details","")),
            esc(c.get("implications","")),
            badge(c.get("criticality","")),
            esc(c.get("lenders_pushback","")),
            esc(c.get("revised_clause","")),
            esc(c.get("rationale",""))
        ]
        rc = row_class(c.get("criticality",""))
        cells = "".join(
            f'<td title="{v}" class="">{v}</td>' if i == 3 else f'<td title="{v}">{v}</td>'
            for i, v in enumerate(vals)
        )
        rows_html += f'<tr class="{rc}">{cells}</tr>'

    header_html = "".join(f"<th>{h}</th>" for h in headers)

    st.markdown(f"""
    <div class="compact-table-wrapper">
      <table class="compact-table">
        <thead><tr>{header_html}</tr></thead>
        <tbody>{rows_html}</tbody>
      </table>
    </div>
    <p style="font-size:11px;color:#888;margin-top:-0.5rem;">
      💡 Hover over any cell to see the full text. Download the Excel for the complete analysis.
    </p>""", unsafe_allow_html=True)

def build_dataframe(clauses):
    import pandas as pd
    return pd.DataFrame([{
        "Clause Name": c.get("clause_name",""),
        "Clause Details": c.get("clause_details",""),
        "Implications": c.get("implications",""),
        "Criticality": c.get("criticality",""),
        "Lender's Likely Pushback": c.get("lenders_pushback",""),
        "Revised Clause": c.get("revised_clause",""),
        "Rationale for Modification": c.get("rationale","")
    } for c in clauses])

def generate_excel_download(df):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Term Sheet Analysis"

        header_fill = PatternFill(start_color="185FA5", end_color="185FA5", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        crit_colours = {"Non-negotiable":"FCEBEB","Can live with it":"FAEEDA","Good to have":"EAF3DE"}

        for col_idx, header in enumerate(df.columns, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.row_dimensions[1].height = 30

        border = Border(
            left=Side(style='thin', color='CCCCCC'), right=Side(style='thin', color='CCCCCC'),
            top=Side(style='thin', color='CCCCCC'), bottom=Side(style='thin', color='CCCCCC')
        )
        for row_idx, row in df.iterrows():
            row_colour = crit_colours.get(row.get('Criticality',''), "FFFFFF")
            row_fill = PatternFill(start_color=row_colour, end_color=row_colour, fill_type="solid")
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx + 2, column=col_idx, value=str(value))
                cell.fill = row_fill
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = border

        for i, width in enumerate([18, 30, 30, 16, 30, 30, 35], 1):
            ws.column_dimensions[get_column_letter(i)].width = width

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output.getvalue(), "xlsx"
    except ImportError:
        return df.to_csv(index=False).encode('utf-8-sig'), "csv"

# ─── Main App ─────────────────────────────────────────────────────────────────
def main():
    # ── Session state initialisation ──
    if "analysis_result" not in st.session_state:
        st.session_state.analysis_result = None
    if "analysis_df" not in st.session_state:
        st.session_state.analysis_df = None
    if "analysis_filename" not in st.session_state:
        st.session_state.analysis_filename = None
    if "excel_data" not in st.session_state:
        st.session_state.excel_data = None
    if "excel_ext" not in st.session_state:
        st.session_state.excel_ext = "xlsx"
    if "usage" not in st.session_state:
        st.session_state.usage = None

    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📄 Term Sheet Scanner</h1>
        <p>AI-powered term sheet analysis for NBFC debt capital markets teams</p>
    </div>""", unsafe_allow_html=True)
    st.divider()

    # ── API Key ──
    api_key = None
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    except (KeyError, FileNotFoundError):
        with st.sidebar:
            st.markdown("### ⚙️ Configuration")
            api_key = st.text_input("Anthropic API Key", type="password", placeholder="sk-ant-...")
            if api_key:
                st.success("API key entered ✓")
            else:
                st.warning("Please enter your API key.")

    # ── Upload Section (shown only when no analysis is stored) ──
    if st.session_state.analysis_result is None:
        st.markdown("### Upload Term Sheet")
        st.caption("Supported formats: PDF, Word (.docx), Excel (.xlsx), Image (JPG, PNG) · Maximum file size: 10MB")

        uploaded_file = st.file_uploader(
            "Drop your term sheet here or click to browse",
            type=["pdf","docx","doc","xlsx","xls","jpg","jpeg","png"],
            label_visibility="collapsed"
        )

        if uploaded_file:
            if uploaded_file.size > 10 * 1024 * 1024:
                st.error("File size exceeds 10MB limit. Please upload a smaller file.")
                st.stop()

            st.success(f"✓ File selected: **{uploaded_file.name}** ({uploaded_file.size/1024:.1f} KB)")

            col1, col2, col3 = st.columns([1,1,1])
            with col2:
                analyse_clicked = st.button("🔍 Analyse Term Sheet", type="primary",
                                            use_container_width=True, disabled=not api_key)

            if not api_key:
                st.info("Please enter your Anthropic API key in the sidebar.")

            if analyse_clicked and api_key:
                with st.spinner("Extracting text from document..."):
                    content, content_type = extract_text(uploaded_file)

                if content_type == "unsupported":
                    st.error("File format not supported.")
                    st.stop()
                if content is None:
                    st.error("Could not extract text. Please ensure the file is not corrupted.")
                    st.stop()
                if content_type == "text" and len(str(content).strip()) < 50:
                    st.error("Could not extract readable text. Try uploading as JPG/PNG if it is a scanned document.")
                    st.stop()

                with st.spinner("Analysing the term sheet — this may take 15–30 seconds..."):
                    result, error, usage = analyse_term_sheet(content, content_type, api_key)

                if error:
                    st.error(f"Analysis error: {error}")
                    st.stop()

                if result.get("not_termsheet"):
                    st.warning("⚠️ Document does not appear to be a term sheet. Please upload the correct document.")
                    st.stop()

                # ── Store results in session state ──
                df = build_dataframe(result.get("clauses", []))
                excel_data, excel_ext = generate_excel_download(df)

                st.session_state.analysis_result = result
                st.session_state.analysis_df = df
                st.session_state.analysis_filename = uploaded_file.name.split('.')[0]
                st.session_state.excel_data = excel_data
                st.session_state.excel_ext = excel_ext
                st.session_state.usage = usage

                st.rerun()

    # ── Results Section (shown when analysis is stored in session state) ──
    if st.session_state.analysis_result is not None:
        result = st.session_state.analysis_result
        df = st.session_state.analysis_df
        clauses = result.get("clauses", [])

        st.markdown("### Analysed Term Sheet")

        if "summary" in result:
            display_summary(result["summary"])

        nn = sum(1 for c in clauses if c.get("criticality") == "Non-negotiable")
        cl = sum(1 for c in clauses if c.get("criticality") == "Can live with it")
        gh = sum(1 for c in clauses if c.get("criticality") == "Good to have")

        col1, col2, col3 = st.columns(3)
        col1.metric("🔴 Non-negotiable", nn)
        col2.metric("🟡 Can live with it", cl)
        col3.metric("🟢 Good to have", gh)

        st.markdown("&nbsp;")
        display_compact_table(clauses)

        st.divider()
        col1, col2 = st.columns([1, 3])

        with col1:
            ext = st.session_state.excel_ext
            st.download_button(
                label="⬇️ Download Analysis",
                data=st.session_state.excel_data,
                file_name=f"term_sheet_analysis_{st.session_state.analysis_filename}.{ext}",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary",
                use_container_width=True
            )

        with col2:
            if st.button("🔄 Analyse Another Term Sheet"):
                st.session_state.analysis_result = None
                st.session_state.analysis_df = None
                st.session_state.analysis_filename = None
                st.session_state.excel_data = None
                st.session_state.excel_ext = "xlsx"
                st.session_state.usage = None
                st.rerun()

        # ── Token Usage & Cost Display ──
        if st.session_state.usage:
            u = st.session_state.usage
            st.markdown("---")
            st.markdown(
                "<p style='font-size:12px;color:#aaa;margin-bottom:4px;'>"
                "📊 API usage for this analysis (Claude Sonnet 4.6 — $3/M input · $15/M output)"
                "</p>",
                unsafe_allow_html=True
            )
            m1, m2, m3, m4, m5 = st.columns(5)
            m1.metric("Input Tokens",  f"{u['input_tokens']:,}")
            m2.metric("Output Tokens", f"{u['output_tokens']:,}")
            m3.metric("Input Cost",    f"${u['input_cost_usd']:.4f}")
            m4.metric("Output Cost",   f"${u['output_cost_usd']:.4f}")
            m5.metric("Total Cost",    f"${u['total_cost_usd']:.4f} · ₹{u['total_cost_inr']:.2f}")

    st.markdown("""
    <div class="disclaimer">
        In MVP-1, analysis is not persisted. If the user closes the browser, the analysis is lost.
        The user must download the file before closing. MVP-1 has no authentication — anyone with the URL can access the tool.
    </div>""", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
